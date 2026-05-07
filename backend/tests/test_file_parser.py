import io

import pytest

from app.services.file_parser import (
    FileParseError,
    parse_docx,
    parse_pdf,
    parse_upload,
    validate_file,
)


def test_validate_pdf_valid():
    validate_file("resume.pdf", b"%PDF-1.4\x00\x00\x00")


def test_validate_docx_valid():
    validate_file("resume.docx", b"PK\x03\x04\x00\x00\x00")


def test_validate_invalid_ext():
    with pytest.raises(FileParseError, match="不支持的文件格式"):
        validate_file("resume.txt", b"hello")


def test_validate_oversized():
    big = b"x" * (6 * 1024 * 1024)
    with pytest.raises(FileParseError, match="超过限制"):
        validate_file("resume.pdf", big)


def test_validate_pdf_bad_magic():
    with pytest.raises(FileParseError, match="不是有效的 PDF"):
        validate_file("resume.pdf", b"\x00\x00\x00\x00\x00")


def test_validate_docx_bad_magic():
    with pytest.raises(FileParseError, match="不是有效的 DOCX"):
        validate_file("resume.docx", b"\x00\x00\x00\x00\x00")


def test_parse_docx_empty():
    from docx import Document

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    text = parse_docx(buf.read())
    assert text == ""


def test_parse_pdf_minimal():
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Hello World")
    c.save()
    buf.seek(0)
    text = parse_pdf(buf.read())
    assert "Hello World" in text


def test_parse_upload_dispatches_pdf():
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "PDF Content")
    c.save()
    buf.seek(0)
    text = parse_upload("resume.pdf", buf.read())
    assert "PDF Content" in text


def test_parse_upload_dispatches_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("DOCX Content")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    text = parse_upload("resume.docx", buf.read())
    assert "DOCX Content" in text


def test_parse_upload_unsupported():
    with pytest.raises(FileParseError, match="不支持的文件格式"):
        parse_upload("resume.txt", b"hello")
