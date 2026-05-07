import io

from docx import Document


def test_upload_docx_returns_parsed_text(client):
    doc = Document()
    doc.add_paragraph("I have experience with Python and FastAPI development.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = client.post(
        "/api/resumes/upload",
        files={"file": ("resume.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"candidate_name": "Alex Chen"},
    )

    assert response.status_code == 201
    data = response.json()
    assert data["candidate_name"] == "Alex Chen"
    assert "Python" in data["raw_text"]
    assert data["parsed_from"] == "resume.docx"
    assert "fastapi" in data["profile"].get("evidence", {})
    assert data["id"] > 0


def test_upload_without_candidate_name_uses_filename(client):
    doc = Document()
    doc.add_paragraph("Python and SQL experience.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = client.post(
        "/api/resumes/upload",
        files={"file": ("my_resume.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 201
    data = response.json()
    assert data["candidate_name"] == "my_resume.docx"
    assert data["version_label"] == "uploaded"


def test_upload_invalid_file_type_returns_422(client):
    response = client.post(
        "/api/resumes/upload",
        files={"file": ("resume.txt", b"hello world", "text/plain")},
    )

    assert response.status_code == 422
    assert "不支持的文件格式" in response.text
